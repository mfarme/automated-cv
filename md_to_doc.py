import os
import markdown
from docx import Document
from bs4 import BeautifulSoup

def convert_to_docx(md_file, save_path):
    try:
        # Read markdown content
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['tables'])
        
        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create docx document
        doc = Document()
        
        # Process each HTML element
        for element in soup.recursiveChildGenerator():
            if element.name:
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    doc.add_heading(element.get_text(), level=level)
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Number')
        
        # Save docx file
        output_file = os.path.join(save_path, md_file.replace('.md', '.docx'))
        doc.save(output_file)
        print(f"DOCX file saved at: {output_file}")
        return output_file
    except Exception as e:
        print(f"Error converting Markdown to DOCX: {str(e)}")
        return None
